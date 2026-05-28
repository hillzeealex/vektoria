"""
Individual format extractors. Each returns a StructuredDocument.
"""

import csv
import io
import re
from html.parser import HTMLParser
from pathlib import Path

from pdf_extractor.models import Section, StructuredDocument


# ── DOCX ─────────────────────────────────────────────────────────────

class DocxExtractor:
    """Extract text from .docx files using python-docx. Detects heading styles."""

    def extract(self, file_path: str) -> StructuredDocument:
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        parts: list[str] = []
        sections: list[Section] = []
        current_section: Section | None = None
        content_buffer: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            heading_level = self._heading_level(para)

            if heading_level:
                if current_section and content_buffer:
                    current_section.content = "\n".join(content_buffer)
                    content_buffer = []

                prefix = "#" * heading_level
                parts.append(f"\n{prefix} {text}\n")
                current_section = Section(
                    title=text, level=heading_level, page=0
                )
                sections.append(current_section)
            else:
                parts.append(text)
                content_buffer.append(text)

        if current_section and content_buffer:
            current_section.content = "\n".join(content_buffer)

        markdown = "\n".join(parts).strip()
        title = sections[0].title if sections else Path(file_path).stem

        return StructuredDocument(
            markdown=markdown,
            page_count=1,
            title=title,
            sections=sections,
            metadata={"parser": "docx-extractor", "source": file_path},
        )

    @staticmethod
    def _heading_level(para) -> int | None:
        style_name = (para.style.name or "").lower()
        if style_name.startswith("heading"):
            # "Heading 1" -> 1, "Heading 2" -> 2, etc.
            parts = style_name.split()
            if len(parts) == 2 and parts[1].isdigit():
                return min(int(parts[1]), 6)
        return None


# ── Markdown / TXT ───────────────────────────────────────────────────

_MD_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")


class MarkdownExtractor:
    """Extract text from .md and .txt files. Detects # headings."""

    def extract(self, file_path: str) -> StructuredDocument:
        text = Path(file_path).read_text(encoding="utf-8")
        lines = text.splitlines()

        parts: list[str] = []
        sections: list[Section] = []
        current_section: Section | None = None
        content_buffer: list[str] = []

        for line in lines:
            m = _MD_HEADING_RE.match(line)
            if m:
                if current_section and content_buffer:
                    current_section.content = "\n".join(content_buffer)
                    content_buffer = []

                level = len(m.group(1))
                title = m.group(2).strip()
                parts.append(f"\n{'#' * level} {title}\n")
                current_section = Section(title=title, level=level, page=0)
                sections.append(current_section)
            else:
                parts.append(line)
                if line.strip():
                    content_buffer.append(line.strip())

        if current_section and content_buffer:
            current_section.content = "\n".join(content_buffer)

        markdown = "\n".join(parts).strip()
        title = sections[0].title if sections else Path(file_path).stem

        return StructuredDocument(
            markdown=markdown,
            page_count=1,
            title=title,
            sections=sections,
            metadata={"parser": "markdown-extractor", "source": file_path},
        )


# ── HTML ─────────────────────────────────────────────────────────────

class _HtmlContentParser(HTMLParser):
    """Minimal HTML parser that extracts headings (h1-h6) and paragraph text."""

    def __init__(self):
        super().__init__()
        self.elements: list[tuple[str, str]] = []  # (tag, text)
        self._current_tag: str | None = None
        self._current_text: list[str] = []
        self._capture_tags = {"h1", "h2", "h3", "h4", "h5", "h6", "p"}

    def handle_starttag(self, tag: str, attrs):
        tag = tag.lower()
        if tag in self._capture_tags:
            self._current_tag = tag
            self._current_text = []

    def handle_endtag(self, tag: str):
        tag = tag.lower()
        if tag == self._current_tag:
            text = "".join(self._current_text).strip()
            if text:
                self.elements.append((tag, text))
            self._current_tag = None
            self._current_text = []

    def handle_data(self, data: str):
        if self._current_tag:
            self._current_text.append(data)


class HtmlExtractor:
    """Extract text from HTML using stdlib html.parser."""

    def extract(self, file_path: str) -> StructuredDocument:
        html_text = Path(file_path).read_text(encoding="utf-8")

        parser = _HtmlContentParser()
        parser.feed(html_text)

        parts: list[str] = []
        sections: list[Section] = []
        current_section: Section | None = None
        content_buffer: list[str] = []

        for tag, text in parser.elements:
            if tag.startswith("h") and tag[1:].isdigit():
                if current_section and content_buffer:
                    current_section.content = "\n".join(content_buffer)
                    content_buffer = []

                level = int(tag[1])
                prefix = "#" * level
                parts.append(f"\n{prefix} {text}\n")
                current_section = Section(title=text, level=level, page=0)
                sections.append(current_section)
            else:
                parts.append(text)
                content_buffer.append(text)

        if current_section and content_buffer:
            current_section.content = "\n".join(content_buffer)

        markdown = "\n".join(parts).strip()
        title = sections[0].title if sections else Path(file_path).stem

        return StructuredDocument(
            markdown=markdown,
            page_count=1,
            title=title,
            sections=sections,
            metadata={"parser": "html-extractor", "source": file_path},
        )


# ── CSV ──────────────────────────────────────────────────────────────

class CsvExtractor:
    """Extract CSV as structured data. Column headers become context."""

    def extract(self, file_path: str) -> StructuredDocument:
        text = Path(file_path).read_text(encoding="utf-8")
        reader = csv.reader(io.StringIO(text))

        rows = list(reader)
        if not rows:
            return StructuredDocument(
                markdown="",
                page_count=1,
                title=Path(file_path).stem,
                sections=[],
                metadata={"parser": "csv-extractor", "source": file_path},
            )

        headers = rows[0]
        data_rows = rows[1:]

        parts: list[str] = []
        sections: list[Section] = []

        # Title section with column info
        title = Path(file_path).stem
        parts.append(f"# {title}\n")
        parts.append(f"Columns: {', '.join(headers)}\n")

        title_section = Section(title=title, level=1, page=0)
        sections.append(title_section)

        # Each row as a content block
        row_texts: list[str] = []
        for i, row in enumerate(data_rows):
            row_parts = []
            for header, value in zip(headers, row):
                if value.strip():
                    row_parts.append(f"**{header}**: {value}")
            if row_parts:
                row_text = " | ".join(row_parts)
                parts.append(f"- {row_text}")
                row_texts.append(row_text)

        title_section.content = "\n".join(row_texts)

        markdown = "\n".join(parts).strip()

        return StructuredDocument(
            markdown=markdown,
            page_count=1,
            title=title,
            sections=sections,
            metadata={
                "parser": "csv-extractor",
                "source": file_path,
                "columns": headers,
                "row_count": len(data_rows),
            },
        )
